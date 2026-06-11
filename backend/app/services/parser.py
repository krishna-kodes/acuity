import re
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument

from app.config import settings
from app.services.ingestion import PageContent, ParsedDocument

# Internal lowercase->uppercase transition signals broken ligature/CID
# extraction (e.g. "aWendance", "OperaJons", "menJoned").
_MIDWORD_CAPS_RE = re.compile(r"[a-z][A-Z]")
_WORD_RE = re.compile(r"[A-Za-z]{2,}")


def _extraction_quality(text: str) -> float:
    """Heuristic [0..1]: fraction of word tokens free of mid-word-caps garbage.

    Returns 1.0 for empty text (nothing to score). Low scores indicate a
    broken font/ToUnicode mapping in the source PDF.
    """
    words = _WORD_RE.findall(text)
    if not words:
        return 1.0
    garbled = sum(1 for w in words if _MIDWORD_CAPS_RE.search(w))
    return 1.0 - (garbled / len(words))


def _extract_pymupdf(path: str) -> list[str]:
    """Per-page text via PyMuPDF (honors the font ToUnicode table better than
    pdfplumber for many PDFs). Returns [] if PyMuPDF is unavailable or the file
    can't be opened — caller falls back to pdfplumber."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return []
    try:
        out: list[str] = []
        with fitz.open(path) as doc:
            for page in doc:
                out.append(page.get_text("text") or "")
        return out
    except Exception:
        return []


def _ocr_page(path: str, page_index: int) -> str | None:
    """OCR a single page as a last resort. No-op (None) if pytesseract or the
    tesseract binary is missing — OCR is an optional enhancement, not a dep."""
    try:
        import fitz
        import pytesseract  # type: ignore[import-untyped]
        from PIL import Image  # type: ignore[import-untyped]
    except ImportError:
        return None
    try:
        import io

        with fitz.open(path) as doc:
            page = doc[page_index]
            pix = page.get_pixmap(dpi=200)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
        return pytesseract.image_to_string(img) or None
    except Exception:
        return None


async def parse_pdf(path: str) -> ParsedDocument:
    pages = []
    pdfplumber_pages: list[str] = []
    fitz_pages = _extract_pymupdf(path)

    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            pp_text = page.extract_text() or ""
            pdfplumber_pages.append(pp_text)

            # Choose the higher-quality extraction per page.
            candidates = [pp_text]
            if i < len(fitz_pages):
                candidates.append(fitz_pages[i])
            text = max(candidates, key=_extraction_quality)

            # OCR fallback only when the best extractor is still garbage.
            if _extraction_quality(text) < settings.extraction_quality_threshold:
                ocr = _ocr_page(path, i)
                if ocr and _extraction_quality(ocr) > _extraction_quality(text):
                    text = ocr

            raw_tables = page.extract_tables() or []
            tables = [
                [[cell or "" for cell in row] for row in table]
                for table in raw_tables
            ]
            pages.append(PageContent(page_number=i + 1, text=text, tables=tables))
    return ParsedDocument(filename=Path(path).name, pages=pages)


_WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


async def parse_docx(path: str) -> ParsedDocument:
    doc = DocxDocument(path)
    text_lines: list[str] = []
    tables: list[list[list[str]]] = []

    for block in doc.element.body:
        tag = block.tag.split("}")[-1]
        if tag == "p":
            # CT_P/CT_R override .text to return all inner text, causing 3× duplication.
            # Only collect w:t leaf nodes which hold the actual string values.
            text = "".join(n.text or "" for n in block.iter(f"{{{_WNS}}}t"))
            if text.strip():
                text_lines.append(text.strip())
        elif tag == "tbl":
            rows = []
            for row in block.iter(f"{{{_WNS}}}tr"):
                cells = [
                    "".join(n.text or "" for n in cell.iter(f"{{{_WNS}}}t"))
                    for cell in row.iter(f"{{{_WNS}}}tc")
                ]
                rows.append(cells)
            if rows:
                tables.append(rows)

    page = PageContent(
        page_number=1,
        text="\n".join(text_lines),
        tables=tables,
    )
    return ParsedDocument(filename=Path(path).name, pages=[page])
