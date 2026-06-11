"""DOCX proposal export with markdown-to-Word conversion."""

from __future__ import annotations

import os
import re
from dataclasses import dataclass, field
from datetime import datetime
from typing import Any

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


@dataclass
class ProposalSection:
    heading: str
    body: str


@dataclass
class ProposalContent:
    title: str
    sections: list[ProposalSection] = field(default_factory=list)


_HEADER_FILL = "2E5FA3"
_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _hex_to_rgb(hex_color: str) -> RGBColor:
    """Convert '#RRGGBB' or 'RRGGBB' to RGBColor."""
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _shade_cell(cell, fill_hex: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _style_table_header(table, primary_hex: str = _HEADER_FILL) -> None:
    """Bold white-on-color header row."""
    for cell in table.rows[0].cells:
        _shade_cell(cell, primary_hex.lstrip("#"))
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = _HEADER_TEXT


def _add_formatted_runs(para, text: str) -> None:
    """Add runs with **bold** inline markers converted to Word bold."""
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            para.add_run(part[2:-2]).bold = True
        elif part:
            para.add_run(part)


# ---------------------------------------------------------------------------
# Markdown-to-Word body renderer
# ---------------------------------------------------------------------------

def _add_markdown_body(doc, text: str, secondary_hex: str | None = None, primary_hex: str = _HEADER_FILL) -> None:
    """Convert LLM markdown output to native Word elements.

    Handles:
    - ## / ### / #### headings (H1/H2 get secondary_hex text color when provided)
    - Lines that are entirely **bold** → sub-heading (Heading 3)
    - - / * / • bullet lists (with 2-space indent detection for List Bullet 2)
    - 1. / 1) numbered lists
    - **inline bold** within any paragraph or bullet
    - Blank lines as paragraph separators
    - Horizontal rules (---, ***, ___) are silently dropped
    """
    if not text or not text.strip():
        return

    lines = text.splitlines()
    para_lines: list[str] = []
    table_rows: list[list[str]] = []

    def _apply_secondary(heading_para) -> None:
        if secondary_hex:
            for run in heading_para.runs:
                run.font.color.rgb = _hex_to_rgb(secondary_hex)

    def _is_table_separator(line: str) -> bool:
        return bool(re.fullmatch(r"[\|\s\-:]+", line.strip()))

    def _parse_table_row(line: str) -> list[str]:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        return cells

    def flush_table() -> None:
        if not table_rows:
            return
        cols = max(len(r) for r in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=cols)
        table.style = "Table Grid"
        for r_idx, row_cells in enumerate(table_rows):
            for c_idx in range(cols):
                cell_text = row_cells[c_idx] if c_idx < len(row_cells) else ""
                cell = table.cell(r_idx, c_idx)
                cell.text = cell_text
                if r_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
        _style_table_header(table, primary_hex if primary_hex else _HEADER_FILL)
        table_rows.clear()

    def flush() -> None:
        flush_table()
        if not para_lines:
            return
        combined = " ".join(ln for ln in para_lines if ln.strip())
        if combined.strip():
            p = doc.add_paragraph()
            _add_formatted_runs(p, combined.strip())
        para_lines.clear()

    for line in lines:
        stripped = line.rstrip()

        if stripped.startswith("## "):
            flush()
            h = doc.add_heading(stripped[3:].strip(), level=2)
            h.paragraph_format.keep_with_next = True
            _apply_secondary(h)

        elif stripped.startswith("### "):
            flush()
            h = doc.add_heading(stripped[4:].strip(), level=3)
            h.paragraph_format.keep_with_next = True

        elif stripped.startswith("#### "):
            flush()
            h = doc.add_heading(stripped[5:].strip(), level=4)
            h.paragraph_format.keep_with_next = True

        # Entire line is **bold text** or **bold text**: → treat as Heading 3
        elif re.fullmatch(r"\*\*[^*]+\*\*:?", stripped.strip()):
            flush()
            heading_text = stripped.strip().lstrip("*").rstrip("*:").strip()
            h = doc.add_heading(heading_text, level=3)
            h.paragraph_format.keep_with_next = True

        # Markdown table row: starts and ends with |
        elif stripped.startswith("|"):
            if _is_table_separator(stripped):
                pass  # skip separator rows (|---|---|)
            else:
                if para_lines:
                    # flush prose before table starts
                    combined = " ".join(ln for ln in para_lines if ln.strip())
                    if combined.strip():
                        p = doc.add_paragraph()
                        _add_formatted_runs(p, combined.strip())
                    para_lines.clear()
                table_rows.append(_parse_table_row(stripped))

        # Bullet: -, *, • (optionally indented)
        elif re.match(r"^(\s*)[-*•] .+", stripped):
            flush()
            indent = len(line) - len(line.lstrip())
            bullet_text = re.sub(r"^[\s]*[-*•]\s+", "", stripped)
            style = "List Bullet 2" if indent >= 2 else "List Bullet"
            p = doc.add_paragraph(style=style)
            _add_formatted_runs(p, bullet_text)

        # Numbered list: 1. or 1)
        elif re.match(r"^(\s*)\d+[.)]\s+.+", stripped):
            flush()
            indent = len(line) - len(line.lstrip())
            item_text = re.sub(r"^\s*\d+[.)]\s+", "", stripped)
            style = "List Number 2" if indent >= 2 else "List Number"
            p = doc.add_paragraph(style=style)
            _add_formatted_runs(p, item_text)

        # Empty line → paragraph break
        elif stripped == "":
            flush()

        # Horizontal rule → drop silently
        elif re.fullmatch(r"[-_*]{3,}", stripped.strip()):
            flush()

        # Regular prose
        else:
            para_lines.append(stripped)

    flush()


# ---------------------------------------------------------------------------
# Document-level chrome
# ---------------------------------------------------------------------------

def _add_cover_page(
    doc,
    title: str,
    company_name: str = "",
    prepared_by: str = "",
    primary_hex: str = _HEADER_FILL,
) -> None:
    """Centered cover page: title with border rule, attribution, date."""
    for _ in range(5):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.font.size = Pt(28)
    run.bold = True

    # Bottom border rule in primary brand color
    pPr = title_para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), primary_hex.lstrip("#"))
    pBdr.append(bottom)
    pPr.append(pBdr)

    doc.add_paragraph()

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("Product Requirements Document")
    sub_run.font.size = Pt(16)

    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
    date_run.font.size = Pt(12)

    # Attribution line if company or PM name provided
    if company_name or prepared_by:
        doc.add_paragraph()
        attr_para = doc.add_paragraph()
        attr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parts = []
        if prepared_by:
            parts.append(f"Prepared by: {prepared_by}")
        if company_name:
            parts.append(company_name)
        attr_run = attr_para.add_run(" · ".join(parts))
        attr_run.font.size = Pt(11)
        attr_run.italic = True

    doc.add_page_break()


def _add_page_numbers(doc) -> None:
    """Add centered 'Page X' footer to every section."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.add_run("Page ")
        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"
        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar_begin)
        run._r.append(instrText)
        run._r.append(fldChar_end)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def generate_proposal_docx(
    project_id: int,
    content: ProposalContent,
    structured_sections: list[dict[str, Any]] | None = None,
    output_dir: str = "documents",
    branding: Any | None = None,
) -> str:
    """Write a structured DOCX proposal. Returns the file path.

    When structured_sections is provided, renders them in ProposalSectionId
    enum order with typed tables for risks, personas, and features. Falls back
    to the legacy ProposalContent path when structured_sections is None.

    When branding is provided (BrandingSettingsResponse), replaces hardcoded
    colors in table headers, H1/H2 heading text, and the cover page.
    """
    os.makedirs(output_dir, exist_ok=True)
    path = os.path.join(output_dir, f"proposal_{project_id}.docx")
    doc = DocxDocument()

    primary_hex = branding.primary_color.lstrip("#") if branding else _HEADER_FILL
    secondary_hex = branding.secondary_color if branding else None
    company_name = branding.company_name if branding else ""
    prepared_by = branding.prepared_by if branding else ""

    _add_cover_page(
        doc,
        content.title,
        company_name=company_name,
        prepared_by=prepared_by,
        primary_hex=primary_hex,
    )
    _add_page_numbers(doc)

    if structured_sections:
        from app.schemas.proposal_sections import ProposalSectionId

        sec_map = {s["section_id"]: s for s in structured_sections}
        for sid in ProposalSectionId:
            sec = sec_map.get(sid.value)
            if not sec:
                continue

            items = sec.get("items") or []
            body = (sec.get("content") or "").strip()

            if not items and not body:
                continue

            h = doc.add_heading(sec["title"], level=1)
            h.paragraph_format.keep_with_next = True
            if secondary_hex:
                for run in h.runs:
                    run.font.color.rgb = _hex_to_rgb(secondary_hex)

            if items and sid.value == "risks_and_mitigations":
                table = doc.add_table(rows=1, cols=2)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                hdr[0].text = "Risk"
                hdr[1].text = "Mitigation"
                _style_table_header(table, primary_hex)
                for item in items:
                    row = table.add_row()
                    row.cells[0].text = str(item.get("risk", ""))
                    row.cells[1].text = str(item.get("mitigation", ""))

            elif items and sid.value == "target_audience":
                table = doc.add_table(rows=1, cols=3)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                hdr[0].text = "Name"
                hdr[1].text = "Role"
                hdr[2].text = "Needs"
                _style_table_header(table, primary_hex)
                for item in items:
                    row = table.add_row()
                    row.cells[0].text = str(item.get("name", ""))
                    row.cells[1].text = str(item.get("role", ""))
                    row.cells[2].text = str(item.get("needs", ""))

            elif items and sid.value == "key_features":
                table = doc.add_table(rows=1, cols=3)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                hdr[0].text = "Scope"
                hdr[1].text = "Feature"
                hdr[2].text = "Description"
                _style_table_header(table, primary_hex)
                for item in items:
                    row = table.add_row()
                    row.cells[0].text = "In" if item.get("in_scope") else "Out"
                    row.cells[1].text = str(item.get("title", ""))
                    row.cells[2].text = str(item.get("description", ""))

            else:
                _add_markdown_body(doc, body, secondary_hex=secondary_hex, primary_hex=primary_hex)

    else:
        for section in content.sections:
            h = doc.add_heading(section.heading, level=1)
            h.paragraph_format.keep_with_next = True
            if secondary_hex:
                for run in h.runs:
                    run.font.color.rgb = _hex_to_rgb(secondary_hex)
            _add_markdown_body(doc, section.body, secondary_hex=secondary_hex)

    doc.save(path)
    return path
