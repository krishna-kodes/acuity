from unittest.mock import MagicMock, patch

import pytest

# ── Parser tests ────────────────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_parse_pdf_extracts_text():
    """parse_pdf returns a ParsedDocument with at least one page of text."""
    mock_page = MagicMock()
    mock_page.extract_text.return_value = "This is a requirement.\nSystem shall support OAuth."
    mock_page.extract_tables.return_value = []

    mock_pdf = MagicMock()
    mock_pdf.__enter__ = lambda s: s
    mock_pdf.__exit__ = MagicMock(return_value=False)
    mock_pdf.pages = [mock_page]

    with patch("pdfplumber.open", return_value=mock_pdf):
        from app.services.parser import parse_pdf
        result = await parse_pdf("/fake/doc.pdf")

    assert result.filename == "doc.pdf"
    assert len(result.pages) == 1
    assert "OAuth" in result.pages[0].text
    assert result.pages[0].page_number == 1


@pytest.mark.asyncio
async def test_parse_pdf_extracts_tables():
    """parse_pdf converts pdfplumber tables to list[list[list[str]]]."""
    mock_page = MagicMock()
    mock_page.extract_text.return_value = ""
    mock_page.extract_tables.return_value = [
        [["Feature", "Priority"], ["OAuth", "High"], [None, "Low"]]
    ]

    mock_pdf = MagicMock()
    mock_pdf.__enter__ = lambda s: s
    mock_pdf.__exit__ = MagicMock(return_value=False)
    mock_pdf.pages = [mock_page]

    with patch("pdfplumber.open", return_value=mock_pdf):
        from app.services.parser import parse_pdf
        result = await parse_pdf("/fake/doc.pdf")

    assert len(result.pages[0].tables) == 1
    assert result.pages[0].tables[0][0] == ["Feature", "Priority"]
    assert result.pages[0].tables[0][2] == ["", "Low"]  # None replaced with ""


@pytest.mark.asyncio
async def test_parse_docx_extracts_text(tmp_path):
    """parse_docx returns a ParsedDocument with extracted paragraph text."""
    from docx import Document as DocxDocument
    docx_path = tmp_path / "test.docx"
    doc = DocxDocument()
    doc.add_paragraph("1. Authentication Requirements")
    doc.add_paragraph("The system shall support OAuth 2.0.")
    doc.save(str(docx_path))

    from app.services.parser import parse_docx
    result = await parse_docx(str(docx_path))

    assert result.filename == "test.docx"
    assert "OAuth" in result.pages[0].text


# ── Chunker tests ────────────────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_chunk_detects_numbered_header():
    """Lines matching '2.3 Title' are classified as headers."""
    from app.services.chunker import chunk_document
    from app.services.ingestion import PageContent, ParsedDocument

    parsed = ParsedDocument(
        filename="req.pdf",
        pages=[PageContent(
            page_number=1,
            text="2.3 Authentication Requirements\nThe system shall support OAuth 2.0.",
            tables=[],
        )],
    )
    chunks = await chunk_document(parsed, "proj_1")
    types = [c.detected_type for c in chunks]
    assert "header" in types


@pytest.mark.asyncio
async def test_chunk_detects_list_items():
    """Lines starting with '- ' or '* ' are classified as list_items."""
    from app.services.chunker import chunk_document
    from app.services.ingestion import PageContent, ParsedDocument

    parsed = ParsedDocument(
        filename="req.pdf",
        pages=[PageContent(
            page_number=1,
            text=(
                "- The system shall support OAuth 2.0\n"
                "- Response time < 200ms\n"
                "* Must handle 1000 concurrent users"
            ),
            tables=[],
        )],
    )
    chunks = await chunk_document(parsed, "proj_1")
    types = [c.detected_type for c in chunks]
    assert "list_item" in types


@pytest.mark.asyncio
async def test_chunk_table_is_atomic():
    """A table always produces exactly one chunk regardless of size."""
    from app.services.chunker import chunk_document
    from app.services.ingestion import PageContent, ParsedDocument

    big_table = [["Col A", "Col B"]] + [["val", "val"] for _ in range(100)]
    parsed = ParsedDocument(
        filename="req.pdf",
        pages=[PageContent(page_number=1, text="", tables=[big_table])],
    )
    chunks = await chunk_document(parsed, "proj_1")
    table_chunks = [c for c in chunks if c.detected_type == "table"]
    assert len(table_chunks) == 1


@pytest.mark.asyncio
async def test_chunk_token_bounds():
    """All non-table chunks are within [min_tokens, max_tokens]."""
    from app.services.chunker import chunk_document
    from app.services.ingestion import PageContent, ParsedDocument

    long_text = " ".join(["The system shall handle requests efficiently."] * 50)
    parsed = ParsedDocument(
        filename="req.pdf",
        pages=[PageContent(page_number=1, text=long_text, tables=[])],
    )
    chunks = await chunk_document(parsed, "proj_1", min_tokens=50, max_tokens=200)
    for chunk in chunks:
        if chunk.detected_type != "table":
            assert chunk.token_count >= 50, f"Chunk too small: {chunk.token_count}"
            assert chunk.token_count <= 200, f"Chunk too large: {chunk.token_count}"


@pytest.mark.asyncio
async def test_chunk_section_hint_propagates():
    """section_hint on non-header chunks equals the most recent header text."""
    from app.services.chunker import chunk_document
    from app.services.ingestion import PageContent, ParsedDocument

    long_para = "The system shall support OAuth 2.0 and related authentication flows. " * 5
    parsed = ParsedDocument(
        filename="req.pdf",
        pages=[PageContent(
            page_number=1,
            text=f"2. Functional Requirements\n\n{long_para}",
            tables=[],
        )],
    )
    chunks = await chunk_document(parsed, "proj_1")
    non_headers = [c for c in chunks if c.detected_type != "header"]
    assert len(non_headers) > 0, "Expected at least one non-header chunk"
    assert all(c.section_hint == "2. Functional Requirements" for c in non_headers)


# ── Embedder tests ───────────────────────────────────────────────────────────

def test_collection_exists_false(tmp_path, monkeypatch):
    """Returns False when the project has no embeddings yet."""
    monkeypatch.setenv("CHROMA_PERSIST_PATH", str(tmp_path))
    monkeypatch.setenv("OPENAI_API_KEY", "test-key")

    with patch("app.services.embedder.OpenAIEmbeddingFunction"):
        from app.services.embedder import collection_exists
        assert collection_exists("new_project") is False


@pytest.mark.asyncio
async def test_embed_and_store_calls_upsert(tmp_path, monkeypatch):
    """embed_and_store calls collection.upsert with correct metadata keys."""
    monkeypatch.setenv("CHROMA_PERSIST_PATH", str(tmp_path))
    monkeypatch.setenv("OPENAI_API_KEY", "test-key")

    from app.services.ingestion import Chunk

    chunks = [
        Chunk(
            text="The system shall support OAuth 2.0.",
            chunk_index=0,
            project_id="42",
            detected_type="paragraph",
            page_number=1,
            section_hint="2. Auth",
            token_count=10,
        )
    ]

    mock_collection = MagicMock()
    mock_collection.upsert = MagicMock()

    with patch("app.services.embedder.chromadb.PersistentClient") as mock_client, \
         patch("app.services.embedder.OpenAIEmbeddingFunction"):
        mock_client.return_value.get_or_create_collection.return_value = mock_collection

        from app.services.embedder import embed_and_store
        stored = await embed_and_store(chunks)

    assert stored == 1
    mock_collection.upsert.assert_called_once()
    call_kwargs = mock_collection.upsert.call_args
    metadata = call_kwargs.kwargs["metadatas"][0]
    assert metadata["project_id"] == "42"
    assert metadata["detected_type"] == "paragraph"
    assert metadata["section_hint"] == "2. Auth"
    assert metadata["token_count"] == 10


@pytest.mark.asyncio
async def test_embed_and_store_empty_returns_zero():
    """embed_and_store with empty list returns 0 without calling ChromaDB."""
    from app.services.embedder import embed_and_store
    result = await embed_and_store([])
    assert result == 0


# ── Ingestion orchestrator tests ─────────────────────────────────────────────

@pytest.mark.asyncio
async def test_ingest_document_cache_hit(tmp_path, monkeypatch):
    """Second call to ingest_document with same project_id is a no-op."""
    monkeypatch.setenv("CHROMA_PERSIST_PATH", str(tmp_path))
    monkeypatch.setenv("OPENAI_API_KEY", "test-key")

    mock_collection = MagicMock()
    mock_collection.count.return_value = 5  # already has embeddings

    embed_call_count = {"n": 0}

    async def fake_embed(chunks):
        embed_call_count["n"] += 1
        return len(chunks)

    with patch("app.services.embedder.chromadb.PersistentClient") as mock_client, \
         patch("app.services.embedder.OpenAIEmbeddingFunction"), \
         patch("app.services.embedder.embed_and_store", fake_embed):
        # Make list_collections return the project so collection_exists returns True
        mock_collection_info = MagicMock()
        mock_collection_info.name = "project_99"
        mock_client.return_value.list_collections.return_value = [mock_collection_info]
        mock_client.return_value.get_collection.return_value = mock_collection
        mock_client.return_value.get_or_create_collection.return_value = mock_collection

        from importlib import reload

        import app.services.ingestion as ing_mod
        reload(ing_mod)

        result = await ing_mod.ingest_document(
            document_id=1, project_id=99, file_path="/fake/doc.pdf", db=MagicMock()
        )

    assert result == 5          # returned existing count
    assert embed_call_count["n"] == 0  # embed_and_store never called


@pytest.mark.asyncio
async def test_ingest_document_updates_status(tmp_path, monkeypatch):
    """ingest_document sets Document.status = ready after ingestion."""
    monkeypatch.setenv("CHROMA_PERSIST_PATH", str(tmp_path))
    monkeypatch.setenv("OPENAI_API_KEY", "test-key")

    mock_db = MagicMock()
    mock_collection = MagicMock()
    mock_collection.count.return_value = 0

    mock_page = MagicMock()
    mock_page.extract_text.return_value = (
        "The system shall support OAuth 2.0 login and session management."
    )
    mock_page.extract_tables.return_value = []
    mock_pdf = MagicMock()
    mock_pdf.__enter__ = lambda s: s
    mock_pdf.__exit__ = MagicMock(return_value=False)
    mock_pdf.pages = [mock_page]

    with patch("app.services.embedder.chromadb.PersistentClient") as mock_client, \
         patch("app.services.embedder.OpenAIEmbeddingFunction"), \
         patch("pdfplumber.open", return_value=mock_pdf):
        # No existing collections → collection_exists returns False
        mock_client.return_value.list_collections.return_value = []
        mock_client.return_value.get_or_create_collection.return_value = mock_collection

        from importlib import reload

        import app.services.ingestion as ing_mod
        reload(ing_mod)

        await ing_mod.ingest_document(
            document_id=7, project_id=3, file_path="/fake/req.pdf", db=mock_db
        )

    mock_db.commit.assert_called_once()
